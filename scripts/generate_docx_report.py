from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # --- FORMATAÇÃO DE ESTILOS ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)

    # Capa Profissional
    doc.add_paragraph("\n\n\n")
    title = doc.add_heading('RELATÓRIO ESTRATÉGICO DE PROJETO', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pardis SellerFlow | Grupo Hermes Pardini & Fleury")
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph("Data de Emissão: 03 de Fevereiro de 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Responsável Técnico: Equipe de Engenharia Antigravity").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 1. VISÃO GERAL E OBJETIVOS ESTRATÉGICOS ---
    doc.add_heading('1. Visão Geral e Objetivos Estratégicos', level=1)
    doc.add_paragraph(
        "A plataforma Pardis SellerFlow foi concebida como o centro nevrálgico da operação de vendas B2B "
        "para o grupo Hermes Pardini e Fleury. O projeto visa mitigar a complexidade operacional da venda de insumos "
        "e serviços laboratoriais, unificando a agilidade do e-commerce (VTEX) com o controle financeiro do ERP (TOTVS)."
    )
    doc.add_paragraph("Objetivos Principais:", style='List Bullet')
    doc.add_paragraph("Eliminar erros de precificação manual em tabelas complexas.", style='List Bullet')
    doc.add_paragraph("Reduzir o ciclo de vida da cotação (Time-to-Quote) em pelo menos 40%.", style='List Bullet')
    doc.add_paragraph("Garantir governança total sobre margens de lucro via fluxo de aprovação automatizado.", style='List Bullet')

    # --- 2. ARQUITETURA DO SISTEMA (CONCLUÍDO) ---
    doc.add_heading('2. Arquitetura e Infraestrutura Técnica', level=1)
    doc.add_paragraph(
        "O sistema utiliza uma arquitetura de micro-serviços moderna baseada em nuvem (Cloud-Native), "
        "distribuída da seguinte forma:"
    )
    doc.add_paragraph(
        "Integração de Catálogo (VTEX Bridge): Sincronização automatizada de mais de 10.000 SKUs via Edge Functions, "
        "garantindo que preços e disponibilidade reflitam exatamente o configurado no Admin VTEX.", style='List Bullet'
    )
    doc.add_paragraph(
        "Banco de Dados de Alta Performance: Utilização de PostgreSQL com Views Materializadas no Supabase, "
        "permitindo buscas de catálogo em milissegundos para os vendedores.", style='List Bullet'
    )
    doc.add_paragraph(
        "Segurança B2B: Camada de autenticação robusta que isola dados sensíveis (margens, custos) do front-end, "
        "expondo apenas o necessário para a operação de campo.", style='List Bullet'
    )

    # --- 3. FUNCIONALIDADES E BENEFÍCIOS ENTREGUES ---
    doc.add_heading('3. Funcionalidades e Benefícios Entregues', level=1)
    
    doc.add_heading('3.1. Engine de Precificação Inteligente', level=2)
    doc.add_paragraph(
        "O sistema gerencia automaticamente a conversão entre unidades laboratoriais e embalagens comerciais, "
        "permitindo descontos agressivos sem perder o rastro da rentabilidade mínima (Floor Price)."
    )

    doc.add_heading('3.2. Novo Fluxo de Negociação Direta', level=2)
    doc.add_paragraph(
        "Implementamos o seletor de produtos dinâmico que permite ao vendedor aplicar 'Price Overrides' "
        "diretamente no carrinho, com cálculo de impostos (ICMS/ST) projetado instantaneamente."
    )

    doc.add_heading('3.3. SLA Baseado em Horário Útil', level=2)
    doc.add_paragraph(
        "Diferente de sistemas comuns, o SellerFlow utiliza um calendário de Business Hours. Se uma "
        "cotação é enviada fora do expediente da Hermes Pardini, o SLA de aprovação congela, reiniciando apenas "
        "no próximo horário comercial. Isso gera métricas de KPI de gestão justas e precisas."
    )

    # --- 4. STATUS ATUAL E INTEGRAÇÃO TOTVS ---
    doc.add_heading('4. Status Atual: Integração ERP TOTVS Datasul', level=1)
    doc.add_paragraph(
        "O sistema está 100% funcional em sua camada de interface, logística de preços e aprovações. "
        "Neste momento, a equipe técnica aguarda a coleta das credenciais de API TOTVS (Client ID, Client Secret e URL de Endpoint)."
    )
    doc.add_paragraph(
        "A importância deste dado: As credenciais permitirão que o SellerFlow 'fale' diretamente com o Datasul "
        "para consultar limites de crédito de clientes reais e enviar pedidos para faturamento automático."
    )

    # --- 5. ROADMAP DE PRÓXIMOS PASSOS ---
    doc.add_heading('5. Roadmap Pós-Coleta de Credenciais TOTVS', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fase'
    hdr_cells[1].text = 'Ação Técnica'
    hdr_cells[2].text = 'Resultados Esperados'

    row = table.add_row().cells
    row[0].text = 'Fase 1'
    row[1].text = 'Homologação de Túnel'
    row[2].text = 'Segurança total na troca de dados entre Supabase e TOTVS.'

    row = table.add_row().cells
    row[0].text = 'Fase 2'
    row[1].text = 'Sync de Clientes'
    row[2].text = 'Fim da digitação manual de novos clientes no portal.'

    row = table.add_row().cells
    row[0].text = 'Fase 3'
    row[1].text = 'Order Bridge'
    row[2].text = 'Cotações aprovadas tornam-se pedidos no ERP em < 5 seg.'

    # --- 6. CONCLUSÃO E ANÁLISE DE IMPACTO ---
    doc.add_heading('6. Conclusão', level=1)
    doc.add_paragraph(
        "Com a finalização dos módulos de SLA e Precificação, o Pardis SellerFlow está pronto para ser a "
        "espinha dorsal da operação comercial. A integração TOTVS é o último passo para transformar a plataforma "
        "em uma ferramenta de automação 360 graus para o Grupo Hermes Pardini & Fleury."
    )

    doc.save('docs/RELATORIO_ESTRATEGICO_DETALHADO_HERMES_PARDINI_FLEURY.docx')
    print("Relatório estratégico ultra-detalhado gerado!")

if __name__ == "__main__":
    create_report()
